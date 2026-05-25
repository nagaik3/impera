#!/usr/bin/env python3
"""
Relatório Mensal v2.0 — Arquivo Morto + Cruzamento RedTrack
Identifica quais tarefas finalizadas no mês têm criativos em teste pelo tráfego.
"""

import calendar
import json
import os
import re
import subprocess
import urllib.request
import sys
from collections import defaultdict
from datetime import datetime, timedelta

sys.path.insert(0, os.path.expanduser("~/Scripts"))
from impera_utils import normalize_person_name

API_TOKEN = os.environ.get("CLICKUP_API_TOKEN", "")
REDTRACK_KEY = os.environ.get("REDTRACK_API_KEY", "")
LIST_COPY = "901324556390"
OUTPUT_DIR = os.path.expanduser("~/Documents")

MESES_NOME = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
}


def cu_fetch_archived():
    tasks, page = [], 0
    while True:
        url = (
            f"https://api.clickup.com/api/v2/list/{LIST_COPY}/task"
            f"?statuses%5B%5D=arquivo+morto&subtasks=true&include_closed=true&page={page}"
        )
        req = urllib.request.Request(url, headers={"Authorization": API_TOKEN})
        with urllib.request.urlopen(req) as resp:
            data = json.loads(resp.read())
        tasks.extend(data.get('tasks', []))
        if data.get('last_page', True):
            break
        page += 1
    return tasks


def rt_campaigns(date_from, date_to):
    """Busca campanhas no RedTrack com referências de criativos."""
    url = (
        f"https://api.redtrack.io/report?api_key={REDTRACK_KEY}"
        f"&group=rt_campaign&date_from={date_from}&date_to={date_to}"
        f"&total=true&per=500"
    )
    with urllib.request.urlopen(url) as resp:
        data = json.loads(resp.read())
    return [c for c in data.get('items', []) if c.get('cost', 0) > 0]


def get_cf(task, field_name):
    for cf in task.get('custom_fields', []):
        if field_name.lower() in cf.get('name', '').lower():
            opts = cf.get('type_config', {}).get('options', [])
            val = cf.get('value')
            if val is not None:
                for o in opts:
                    if o.get('orderindex') == val:
                        return normalize_person_name(o['name']) or '—'
    return '—'


def extract_task_refs(name):
    """Extrai todas as referências (AD numbers, creative IDs) de uma tarefa do ClickUp."""
    refs = set()
    upper = name.upper()

    # Pattern: [ADC88], [ADC88V2], [ADC71V12], [ADCE39]
    for m in re.findall(r'\[(ADC?E?\d+(?:V\d+)?)\]', upper):
        refs.add(m)
        # Also add without AD prefix: ADC88 -> C88
        clean = re.sub(r'^ADC?E?', 'C', m) if m.startswith('AD') else m
        refs.add(clean)

    # Pattern: [IMG 644], [IMG 644 V9]
    for m in re.findall(r'IMG\s*(\d+(?:\s*V\d+)?)', upper):
        refs.add(f"IMG{m.replace(' ', '')}")
        refs.add(f"IMG {m.strip()}")
        # Also just the number
        num_only = re.match(r'(\d+)', m).group(1)
        refs.add(num_only)

    # Pattern: AD ranges like [AD75-AD80] or [AD75 ao AD80]
    range_match = re.search(r'AD\s*(\d+)\s*(?:ao|-)\s*AD\s*(\d+)', upper)
    if range_match:
        low = int(range_match.group(1))
        high = int(range_match.group(2))
        for n in range(low, high + 1):
            refs.add(f"AD{n}")
            refs.add(f"AD {n}")
            refs.add(str(n).zfill(2))

    # Single AD pattern: [AD40] or AD40
    for m in re.findall(r'\bAD\s*(\d+)\b', upper):
        if not range_match:  # avoid double-counting from range
            refs.add(f"AD{m}")
            refs.add(f"AD {m}")

    return refs


def extract_campaign_refs(rt_name):
    """Extrai referências de criativos do nome de uma rt_campaign."""
    refs = set()
    upper = rt_name.upper()

    # AD C88, AD CE31, AD C71 V12
    for m in re.findall(r'AD\s+(C[A-Z]?\d+(?:\s*V\d+)?)', upper):
        clean = m.replace(' ', '')
        refs.add(f"AD{clean}")
        refs.add(clean)

    # AD 644, AD 644 V9
    for m in re.findall(r'AD\s+(\d+(?:\s*V\d+)?)', upper):
        clean = m.replace(' ', '')
        refs.add(f"AD{clean}")
        refs.add(clean)
        num = re.match(r'(\d+)', clean).group(1)
        refs.add(num)
        refs.add(f"IMG{num}")

    # AD ranges: AD 140/144, AD140-144, AD145-149
    for m in re.finditer(r'AD\s*(\d+)\s*[/\-]\s*(?:AD\s*)?(\d+)', upper):
        low, high = int(m.group(1)), int(m.group(2))
        for n in range(low, high + 1):
            refs.add(f"AD{n}")
            refs.add(str(n).zfill(2))

    # Multiple ADs separated by commas: "AD CE31, CE38, CE39"
    cesection = re.search(r'AD\s+([\w\s,V\d]+?)(?:\s+-|\s+CBO|$)', upper)
    if cesection:
        items = re.split(r'[,\s]+', cesection.group(1))
        for item in items:
            item = item.strip()
            if re.match(r'^C[A-Z]?\d+', item) or re.match(r'^\d+', item):
                refs.add(item)
                refs.add(f"AD{item}")

    # Specific patterns: V187 H7, etc
    for m in re.findall(r'V(\d+)\s*[HV]\d*', upper):
        refs.add(f"V{m}")

    return refs


def task_in_test(task_refs, all_campaign_refs):
    """Verifica se alguma referência da tarefa está em alguma campanha em teste."""
    matches = task_refs & all_campaign_refs
    return len(matches) > 0, matches


def get_target_month():
    today = datetime.now()
    if today.month == 1:
        return today.year - 1, 12
    return today.year, today.month - 1


def filter_by_month(tasks, year, month):
    month_start = datetime(year, month, 1, 0, 0, 0).timestamp() * 1000
    last_day = calendar.monthrange(year, month)[1]
    month_end = datetime(year, month, last_day, 23, 59, 59).timestamp() * 1000
    return [t for t in tasks if t.get('date_done') and month_start <= int(t['date_done']) <= month_end]


def build_docx(tasks, year, month, in_test_map):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    mes_nome = MESES_NOME[month]
    title = doc.add_heading(
        f"Relatório Mensal v2.0 — Arquivo Morto + Tráfego | {mes_nome} de {year}",
        level=1,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Stats
    em_teste = sum(1 for t in tasks if in_test_map.get(t['id'], False))
    nao_teste = len(tasks) - em_teste

    doc.add_paragraph(
        f"Total: {len(tasks)} tarefas finalizadas em {mes_nome}/{year} | "
        f"Em teste no tráfego: {em_teste} ({em_teste*100//len(tasks) if tasks else 0}%) | "
        f"Não em teste: {nao_teste} | "
        f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )

    # Aggregations
    por_copy = defaultdict(lambda: {'total': 0, 'em_teste': 0})
    por_editor = defaultdict(lambda: {'total': 0, 'em_teste': 0})
    for t in tasks:
        cw = get_cf(t, 'Copywritter')
        ed = get_cf(t, 'Editor de Video')
        in_t = in_test_map.get(t['id'], False)
        por_copy[cw]['total'] += 1
        por_editor[ed]['total'] += 1
        if in_t:
            por_copy[cw]['em_teste'] += 1
            por_editor[ed]['em_teste'] += 1

    # Summary
    doc.add_heading("Resumo por Copywriter", level=2)
    headers = ["Copywriter", "Total Tarefas", "Em Teste", "% Em Teste"]
    table = doc.add_table(rows=1 + len(por_copy), cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, (k, v) in enumerate(sorted(por_copy.items(), key=lambda x: -x[1]['total'])):
        pct = (v['em_teste'] * 100 // v['total']) if v['total'] > 0 else 0
        for ci, val in enumerate([k, str(v['total']), str(v['em_teste']), f"{pct}%"]):
            table.rows[ri + 1].cells[ci].text = val
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph("")
    doc.add_heading("Resumo por Editor de Vídeo", level=2)
    table = doc.add_table(rows=1 + len(por_editor), cols=4)
    table.style = "Light Grid Accent 1"
    headers_e = ["Editor", "Total Tarefas", "Em Teste", "% Em Teste"]
    for i, h in enumerate(headers_e):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, (k, v) in enumerate(sorted(por_editor.items(), key=lambda x: -x[1]['total'])):
        pct = (v['em_teste'] * 100 // v['total']) if v['total'] > 0 else 0
        for ci, val in enumerate([k, str(v['total']), str(v['em_teste']), f"{pct}%"]):
            table.rows[ri + 1].cells[ci].text = val
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph("")

    # Main table
    doc.add_heading("Tarefas Detalhadas (ordenadas por nomenclatura)", level=2)

    headers_t = ["Tarefa", "Copywriter", "Editor", "Conclusão", "Em Teste?"]
    sorted_tasks = sorted(tasks, key=lambda t: t['name'])
    table = doc.add_table(rows=1 + len(sorted_tasks), cols=len(headers_t))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers_t):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    for ri, t in enumerate(sorted_tasks):
        copy = get_cf(t, 'Copywritter')
        editor = get_cf(t, 'Editor de Video')
        dd = datetime.fromtimestamp(int(t['date_done']) / 1000).strftime('%d/%m/%Y')
        in_t = in_test_map.get(t['id'], False)
        teste_label = "✅ SIM" if in_t else "—"

        for ci, val in enumerate([t['name'], copy, editor, dd, teste_label]):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(9)
                    if ci == 4 and in_t:
                        r.font.color.rgb = RGBColor(0x10, 0xa5, 0x6c)
                        r.bold = True

    return doc


def main():
    if not API_TOKEN or not REDTRACK_KEY:
        print("ERRO: tokens não configurados")
        return

    year, month = get_target_month()
    mes_nome = MESES_NOME[month]
    print(f"Gerando relatório v2.0 para {mes_nome}/{year}...")

    print("  Buscando ClickUp (arquivo morto)...")
    all_tasks = cu_fetch_archived()
    filtered = filter_by_month(all_tasks, year, month)
    print(f"    {len(filtered)} tarefas finalizadas no mês")

    # RedTrack: buscar campanhas com cost no mês + 30 dias depois (criativos podem entrar em teste após)
    rt_start = f"{year}-{month:02d}-01"
    next_month = month + 1 if month < 12 else 1
    next_year = year if month < 12 else year + 1
    last_day_next = calendar.monthrange(next_year, next_month)[1]
    rt_end = f"{next_year}-{next_month:02d}-{last_day_next:02d}"

    print(f"  Buscando RedTrack ({rt_start} a {rt_end})...")
    rt_camps = rt_campaigns(rt_start, rt_end)
    print(f"    {len(rt_camps)} campanhas com cost > 0")

    # Build set of all campaign refs
    all_campaign_refs = set()
    for c in rt_camps:
        all_campaign_refs.update(extract_campaign_refs(c.get('rt_campaign', '')))
    print(f"    {len(all_campaign_refs)} referências únicas extraídas")

    # Match each task
    in_test_map = {}
    for t in filtered:
        task_refs = extract_task_refs(t['name'])
        in_test, matches = task_in_test(task_refs, all_campaign_refs)
        in_test_map[t['id']] = in_test

    em_teste = sum(1 for v in in_test_map.values() if v)
    print(f"  {em_teste}/{len(filtered)} tarefas com criativos em teste")

    print("Gerando .docx...")
    doc = build_docx(filtered, year, month, in_test_map)

    filename = f"Relatorio_Mensal_v2_{mes_nome}_{year}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    print(f"Salvo: {output_path}")

    subprocess.run(["open", output_path])
    subprocess.run([
        "osascript", "-e",
        f'display notification "Relatório v2.0 {mes_nome}/{year} gerado!" '
        f'with title "IMPERA - Arquivo Morto + Tráfego" subtitle "{em_teste} em teste"',
    ])
    print("Concluído!")


if __name__ == "__main__":
    main()
