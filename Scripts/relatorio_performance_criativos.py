#!/usr/bin/env python3
"""
Relatório de Performance por Criativo/Copywriter — IMPERA
Cruza RedTrack (performance) + ClickUp (atribuição copywriter).

Uso:
  python3 relatorio_performance_criativos.py                    # Semana anterior (seg-dom)
  python3 relatorio_performance_criativos.py 01/04 21/04        # Período customizado DD/MM
  python3 relatorio_performance_criativos.py 01/04/2026 21/04/2026  # Com ano
  python3 relatorio_performance_criativos.py --notify            # Envia .docx pelo Telegram
"""

import os, sys, re, json, time as _time
import urllib.request
from datetime import datetime, timedelta
from collections import defaultdict

sys.path.insert(0, os.path.expanduser("~/Scripts"))
from impera_utils import detect_nicho
from impera_cache import cached_rt_adgroups, cached_cu_tasks

REDTRACK_KEY = os.environ.get("REDTRACK_API_KEY", "")
CLICKUP_TOKEN = os.environ.get("CLICKUP_API_TOKEN", "")
TELEGRAM_TOKEN = os.environ.get("TELEGRAM_GPDR_TOKEN", "")
TELEGRAM_CHAT_ID = os.environ.get("TELEGRAM_CHAT_ID", "")

COPY_LIST = "901324556390"
# Lista unificada — EDIÇÃO agora na COPY
TRAFEGO_LIST = "901324476398"

NICHO_KW = {
    "EMAGRECIMENTO": "EM", "DIABETES": "DB", "NEUROPATIA": "NE",
    "ADULTO": "ED", "MEMORIA": "MM", "PROSTATA": "PT",
    "ZUMBIDO": "ZB", "ARTICULAR": "DA", "DORES": "DA", "VISAO": "VS",
}
NICHO_FULL = {
    "EM": "Emagrecimento", "DB": "Diabetes", "NE": "Neuropatia",
    "ED": "Adulto", "MM": "Memória", "PT": "Próstata",
    "DA": "Dores Articulares", "ZB": "Zumbido", "VS": "Visão",
}
NICHO_ALIASES = {
    "NE": ["NE", "NEU", "??"], "EM": ["EM"], "DB": ["DB"], "ED": ["ED"],
    "MM": ["MM"], "PT": ["PT"], "DA": ["DA"], "ZB": ["ZB"], "VS": ["VS"],
}
RIP_COPY = {"CE": "ELIAS", "CY": "YAN", "CC": "CASSIO"}
MANUAL_YAN = {"AD986", "AD1022", "AD740", "AD995"}

# Assignee name → copywriter name (for when dropdown is empty)
ASSIGNEE_TO_COPY = {
    "reaper": "REAPER", "cassio": "REAPER",
    "yan": "YAN", "yan da silva": "YAN", "yan da silva rangel": "YAN",
    "crispim": "CRISPIM", "crispim.copywriter": "CRISPIM",
    "ana paula": "ANA", "ana": "ANA",
    "elias": "ELIAS",
}
# Known editors (skip as copywriter candidates)
EDITORS_SET = {
    "wells", "wells lima", "igor oliveira", "igor paiva", "nicolas",
    "muryllo", "editores impera", "lucas - grego ads", "lucas grego",
}


# ============================================================
# HELPERS
# ============================================================

def rt_fetch(params):
    url = "https://api.redtrack.io/report?" + "&".join(f"{k}={v}" for k, v in params.items())
    with urllib.request.urlopen(urllib.request.Request(url)) as resp:
        return json.loads(resp.read())


def parse_camp(name):
    upper = name.upper()
    nicho = None
    for kw, code in NICHO_KW.items():
        if kw in upper:
            nicho = code
            break
    fonte = None
    fm = re.search(r"\[(FB|GG|YT|TT|KW|MG|TB|OB)\]", upper)
    if fm:
        fonte = fm.group(1)
    return nicho, fonte


def clean_ag_name(name):
    n = name.strip()
    n = re.sub(r"\s*[—-]\s*[Cc][oó]pia.*$", "", n)
    n = re.sub(r"\s*-\s*Copy.*$", "", n, flags=re.IGNORECASE)
    n = re.sub(r"\s*[-—]\s*\d{2}/\d{2}.*$", "", n)
    n = re.sub(r"\s*[-—]\s*\d{4}.*$", "", n)
    n = re.sub(r"^\d+[a-zA-Z]\s+G?\d+\s+\d+V\s+[\d-]+\s+", "", n)
    n = re.sub(r"^\d+[a-z]\d+v\d+\s+", "", n, flags=re.IGNORECASE)
    n = re.sub(r"^0\d\s+(?!V)", "", n)
    return re.sub(r"\s+", " ", n).strip()


def extract_creative_id(clean_name):
    """
    Extrai base_id e version de um nome de criativo RT.
    Dois níveis:
      - 'AD76 V10'       → base=AD76, version=V10 (buscar em ranges de AD76)
      - 'AD76 V10 - V9'  → base=AD76V10, version=V9 (buscar em ranges de AD76V10)
      - 'C123'           → base=C123, version=None (ripagem Douglas)
      - 'C123 V3'        → base=C123, version=V3 (variação, buscar copywriter)
    """
    upper = clean_name.upper().strip()

    # Pre-clean: normalize "AD C##" → "ADC##" (RT writes "AD C71 V12", CU uses "ADC71V12")
    # Keeps ADC prefix for CE/CY/CC (AD CE34 → CE34, AD CY05 → CY05)
    ad_c_strip = re.match(r"^AD\s+(C(?:E|Y|C)\s*\d.*)$", upper)
    if ad_c_strip:
        # CE/CY/CC: strip AD prefix (CU uses CE34, not ADCE34)
        upper = ad_c_strip.group(1).strip()
    else:
        # AD C## → ADC## (CU uses ADC71V12, ADC88V2)
        ad_c_merge = re.match(r"^AD\s*C(\d.*)$", upper)
        if ad_c_merge:
            upper = "ADC" + ad_c_merge.group(1).strip()

    # Step 1: Detect two-version pattern (base V## - V##)
    # "AD 76 V10 - V9", "644 V9 - V164", "C71 V12 - V22"
    two_ver = re.match(
        r"(?:AD\s*)?(\d+)\s*V\s*(\d+)\s*[-—]\s*V\s*(\d+)",
        upper
    )
    if not two_ver:
        # Also match C## V## - V## and ADC## V## - V## (C71 V12 - V22, ADC71 V12 - V44)
        two_ver_c = re.match(
            r"((?:ADC|C)\d+)\s*V\s*(\d+)\s*[-—]\s*V\s*(\d+)",
            upper
        )
        if two_ver_c:
            base_id = f"{two_ver_c.group(1)}V{int(two_ver_c.group(2))}"
            return base_id, f"V{int(two_ver_c.group(3))}"
    if two_ver:
        num = int(two_ver.group(1))
        v_base = int(two_ver.group(2))
        v_variation = int(two_ver.group(3))
        base_id = f"AD{num}V{v_base}"
        # For CE/CY/CC prefix
        for px in ["CE", "CY", "CC"]:
            if upper.startswith(px):
                base_id = f"{px}{num}V{v_base}"
                break
        return base_id, f"V{v_variation}"

    # Also handle: "CE34 V13", "CE34 - V17" (single version with CE prefix)
    # These are ripagem variations, version IS the variation
    for px in ["CE", "CY", "CC"]:
        m = re.match(rf"^{px}\s*(\d+)(?:\s*[-—]\s*V\s*(\d+)|\s+V\s*(\d+))?", upper)
        if m:
            base_id = f"{px}{int(m.group(1))}"
            v = m.group(2) or m.group(3)
            version = f"V{int(v)}" if v else None
            return base_id, version

    # Step 2: Single version pattern
    # "AD76 V10" → base=AD76, version=V10
    # "76 V10" → base=AD76, version=V10
    # "C123 V3" → base=C123, version=V3
    base_id = None
    version = None

    # ADC## pattern (e.g. ADC88 V290) — must come before C## and AD##
    adc_m = re.match(r"^ADC(\d+)", upper)
    if adc_m:
        base_id = f"ADC{int(adc_m.group(1))}"

    # C## pattern (ripagem) — C1, C63, C123 etc.
    if not base_id:
        cm = re.match(r"^C(\d+)(?:\b|V|$)", upper)
        if cm:
            base_id = f"C{int(cm.group(1))}"

    # AD## pattern
    if not base_id:
        m = re.search(r"AD\s*(\d+)", upper)
        if m:
            base_id = f"AD{int(m.group(1))}"

    # Bare number: "232V34", "644 V9" — number followed by V or space or end
    if not base_id:
        m = re.match(r"^(\d+)(?:\b|V|$)", upper)
        if m:
            base_id = f"AD{int(m.group(1))}"

    # Version: take the LAST V## in the name
    # \b doesn't work between digit and V (232V34), so also match digit-V boundary
    all_versions = re.findall(r"(?:\b|(?<=\d))V\s*(\d+)", upper)
    if all_versions:
        version = f"V{int(all_versions[-1])}"

    return base_id, version


def fetch_cu_tasks(list_id):
    tasks = []
    page = 0
    while True:
        url = f"https://api.clickup.com/api/v2/list/{list_id}/task?page={page}&limit=100&subtasks=true&include_closed=true"
        req = urllib.request.Request(url)
        req.add_header("Authorization", CLICKUP_TOKEN)
        with urllib.request.urlopen(req) as resp:
            data = json.loads(resp.read())
        batch = data.get("tasks", [])
        tasks.extend(batch)
        if not batch or data.get("last_page", True):
            break
        page += 1
    return tasks


def get_cf_dropdown(task, search_term):
    for cf in task.get("custom_fields", []):
        cf_name = re.sub(r"[^\w\s]", "", cf.get("name", "")).strip().lower()
        if search_term.lower() in cf_name:
            opts = cf.get("type_config", {}).get("options", [])
            val = cf.get("value")
            if val is not None:
                for o in opts:
                    if o.get("orderindex") == val:
                        return o["name"].upper().strip()
    return None


def get_copywriter_from_assignees(task):
    """Fallback: identifica copywriter pelo assignee da tarefa (pulando editores)."""
    for a in task.get("assignees", []):
        username = (a.get("username") or a.get("email") or "").lower().strip()
        if not username:
            continue
        # Skip editors
        if any(ed in username for ed in EDITORS_SET):
            continue
        # Match known copywriters
        for key, cw in ASSIGNEE_TO_COPY.items():
            if key in username:
                return cw
    return None


def parse_cu_task(task):
    name = task["name"]
    nicho = detect_nicho(name)
    copywriter = get_cf_dropdown(task, "copywritter") or get_cf_dropdown(task, "copy") or None
    # Fallback: use assignee if dropdown is empty
    if not copywriter or copywriter == "N/A":
        copywriter = get_copywriter_from_assignees(task) or "N/A"
    editor = get_cf_dropdown(task, "editor") or "N/A"
    all_ids = []
    # AD range: [AD192-AD197]
    ar = re.search(r"AD\s*(\d+)\s*[-aA]\s*(?:AD)?\s*(\d+)", name, re.IGNORECASE)
    # Single AD in brackets: [AD76]
    a1 = re.search(r"\[AD(\d+)\]", name)
    # AD##V## format: [AD76V10], [ADC88V2] — two-level creative
    av = re.search(r"\[AD(\d+)V(\d+)\]", name, re.IGNORECASE)
    # Standalone AD: "AD644V9 - V164"
    a_standalone = re.search(r"(?:^|\s)AD\s*(\d+)(?:V(\d+))?", name, re.IGNORECASE)
    if ar:
        for i in range(int(ar.group(1)), int(ar.group(2)) + 1):
            all_ids.append(f"AD{i}")
    elif av:
        # Two-level: [AD76V10] → register both AD76 and AD76V10 as IDs
        all_ids.append(f"AD{int(av.group(1))}")
        all_ids.append(f"AD{int(av.group(1))}V{int(av.group(2))}")
    # ADC## format: [ADC71], [ADC71V12] — the "AD" is prefix noise, real ref is C71
    adc = re.search(r"\[ADC(\d+)(?:V(\d+))?\]", name, re.IGNORECASE)
    if adc:
        c_num = int(adc.group(1))
        all_ids.append(f"C{c_num}")
        all_ids.append(f"ADC{c_num}")  # Also register the ADC form
        if adc.group(2):
            v_num = int(adc.group(2))
            all_ids.append(f"C{c_num}V{v_num}")
            all_ids.append(f"ADC{c_num}V{v_num}")
    elif a1:
        all_ids.append(f"AD{int(a1.group(1))}")
    elif a_standalone:
        all_ids.append(f"AD{int(a_standalone.group(1))}")
        if a_standalone.group(2):
            all_ids.append(f"AD{int(a_standalone.group(1))}V{int(a_standalone.group(2))}")
    im = re.search(r"\[IMG\s*(\d+)", name, re.IGNORECASE)
    if im:
        all_ids.append(f"AD{int(im.group(1))}")
    for px in ["CE", "CY", "CC"]:
        rr = re.search(rf"\[{px}(\d+)\s*-\s*{px}?(\d+)\]", name, re.IGNORECASE)
        rs = re.search(rf"\[{px}(\d+)\]", name, re.IGNORECASE)
        if rr:
            for i in range(int(rr.group(1)), int(rr.group(2)) + 1):
                all_ids.append(f"{px}{i}")
        elif rs:
            all_ids.append(f"{px}{int(rs.group(1))}")
    cr = re.search(r"\[C(\d+)\s*-\s*C?(\d+)\]", name)
    cs = re.search(r"\[C(\d+)\]", name)
    if cr:
        for i in range(int(cr.group(1)), int(cr.group(2)) + 1):
            all_ids.append(f"C{i}")
    elif cs:
        all_ids.append(f"C{int(cs.group(1))}")
    vr = re.search(r"\[V(\d+)\s*-\s*V?(\d+)\]", name, re.IGNORECASE)
    vs_m = re.search(r"\[V(\d+)\]", name)
    versions = None
    if vr:
        versions = (int(vr.group(1)), int(vr.group(2)))
    elif vs_m:
        versions = (int(vs_m.group(1)), int(vs_m.group(1)))
    return {
        "name": name, "nicho": nicho, "copywriter": copywriter,
        "editor": editor, "all_ids": all_ids, "versions": versions,
    }


def nicho_matches(a, b):
    if a == b:
        return True
    return b in NICHO_ALIASES.get(a, [a])


def find_copywriter(base_id, version, nicho, tasks):
    if base_id in MANUAL_YAN:
        return "YAN", "?", "MANUAL"
    ver_num = None
    if version and version != "-":
        vm = re.search(r"(\d+)", version)
        if vm:
            ver_num = int(vm.group(1))

    # Step 1: Find candidates matching nicho + base_id
    # For two-level base (AD76V10, C71V12): search ONLY tasks that have the exact two-level ID
    # For single-level base (AD76, C71): search tasks with the base ID but EXCLUDE two-level tasks
    is_two_level = bool(re.match(r"^(?:AD|C|CE|CY|CC)\d+V\d+$", base_id))

    if is_two_level:
        # Exact match for two-level: only tasks containing AD76V10 / C71V12
        candidates = [t for t in tasks if nicho_matches(nicho, t["nicho"]) and base_id in t["all_ids"]]
        if not candidates:
            candidates = [t for t in tasks if base_id in t["all_ids"]]
        # Fallback to single-level if nothing found
        if not candidates:
            fallback_base = re.match(r"^((?:AD|C|CE|CY|CC)\d+)", base_id).group(1)
            candidates = [t for t in tasks if nicho_matches(nicho, t["nicho"]) and fallback_base in t["all_ids"]]
    else:
        # Single-level: search for base_id but prefer tasks that DON'T have a V in their base
        # This prevents C71 from matching [ADC71V12][V09-V12] when looking for C71 V12 (level 1)
        all_candidates = [t for t in tasks if nicho_matches(nicho, t["nicho"]) and base_id in t["all_ids"]]
        if not all_candidates:
            all_candidates = [t for t in tasks if base_id in t["all_ids"]]
        # Filter: prefer tasks whose name contains [base_id] directly, not [base_idV##]
        # i.e., for C71, prefer [ADC71][V2-13] over [ADC71V12][V09-V12]
        level1_only = [t for t in all_candidates if not any(
            re.match(rf"^{re.escape(base_id)}V\d+$", aid) for aid in t["all_ids"]
            if aid != base_id
        )]
        candidates = level1_only if level1_only else all_candidates

    # Ripagem without version = credit the ripper, not the variation writer
    # C## = Douglas, CE## = Elias, CY## = Yan, CC## = Cassio
    if not version or version == "-":
        if re.match(r"^C\d+$", base_id):
            return "DOUGLAS*", "?", "C##_INFERRED"
        for px, cp in RIP_COPY.items():
            if base_id.startswith(px) and re.match(rf"^{px}\d+$", base_id):
                return cp, "?", "PREFIXO_BASE"

    if not candidates:
        # Step 2: Fallback by prefix
        for px, cp in RIP_COPY.items():
            if base_id.startswith(px):
                return cp, "?", "PREFIXO"
        if re.match(r"^C\d+$", base_id):
            return "DOUGLAS*", "?", "C##_INFERRED"
        return "SEM ATRIBUIÇÃO", "?", "NONE"

    # Step 3: Filter out N/A copywriters — prefer candidates with actual attribution
    attributed = [t for t in candidates if t["copywriter"] not in ("N/A", None, "")]
    pool = attributed if attributed else candidates

    # Step 4: Version matching (most precise)
    if ver_num is not None:
        for t in pool:
            if t["versions"] and t["versions"][0] <= ver_num <= t["versions"][1]:
                return t["copywriter"], t["editor"], "VER_MATCH"
        # Also check unattributed if attributed didn't match version
        if attributed:
            for t in candidates:
                if t["versions"] and t["versions"][0] <= ver_num <= t["versions"][1]:
                    if t["copywriter"] not in ("N/A", None, ""):
                        return t["copywriter"], t["editor"], "VER_MATCH"

    # Step 5: Single match
    if len(pool) == 1:
        return pool[0]["copywriter"], pool[0]["editor"], "MATCH"

    # Step 6: Multiple candidates — prefer the one with most specific version range
    if ver_num is not None:
        closest = None
        closest_dist = float("inf")
        for t in pool:
            if t["versions"]:
                mid = (t["versions"][0] + t["versions"][1]) / 2
                dist = abs(ver_num - mid)
                if dist < closest_dist:
                    closest = t
                    closest_dist = dist
        if closest:
            return closest["copywriter"], closest["editor"], "CLOSEST_VER"

    return pool[0]["copywriter"], pool[0]["editor"], "FIRST"


# ============================================================
# DATA COLLECTION
# ============================================================

def collect_data(date_from, date_to, progress_fn=None):
    """Coleta e processa todos os dados. Retorna lista de criativos."""

    def log(msg):
        if progress_fn:
            progress_fn(msg)
        print(msg)

    log("Buscando dados RedTrack...")
    rt_data = cached_rt_adgroups(date_from, date_to, ttl=1800)
    campaigns_raw = rt_data["campaigns"]
    all_ags_raw = rt_data["adgroups"]
    active_camps = [c for c in campaigns_raw if float(c.get("cost", 0)) > 0]
    log(f"  {len(active_camps)} campanhas ativas, {len(all_ags_raw)} adgroups")

    all_adgroups = []
    for ag in all_ags_raw:
        cost = float(ag.get("cost", 0))
        if cost < 1:
            continue
        cname = ag.get("campaign", "")
        nicho, fonte = parse_camp(cname)
        fr = float(ag.get("revenuetype2", 0)) + float(ag.get("revenuetype3", 0))
        all_adgroups.append({
            "name": ag.get("rt_adgroup", ""),
            "campaign": cname, "nicho": nicho, "fonte": fonte,
            "cost": cost, "front_rev": fr,
            "total_rev": float(ag.get("total_revenue", 0)),
            "mc_br": fr * 0.74 - cost * 1.12,
            "vendas": int(ag.get("convtype1", 0)),
        })
    log(f"  {len(all_adgroups)} adgroups coletados")

    # Aggregate
    creative_agg = defaultdict(lambda: {
        "cost": 0, "front_rev": 0, "total_rev": 0, "mc_br": 0,
        "vendas": 0, "nicho": None, "fonte": None,
    })
    for ag in all_adgroups:
        clean = clean_ag_name(ag["name"])
        base_id, version = extract_creative_id(clean)
        if not base_id:
            continue
        key = (base_id, version or "-", ag["nicho"] or "?")
        d = creative_agg[key]
        d["cost"] += ag["cost"]
        d["front_rev"] += ag["front_rev"]
        d["total_rev"] += ag["total_rev"]
        d["mc_br"] += ag["mc_br"]
        d["vendas"] += ag["vendas"]
        d["nicho"] = ag["nicho"]
        d["fonte"] = ag["fonte"]

    # ClickUp — COPY + EDIÇÃO + TRÁFEGO (cruzamento triplo)
    log("Buscando tarefas ClickUp (COPY + EDIÇÃO + TRÁFEGO)...")
    copy_raw = cached_cu_tasks(COPY_LIST, include_closed=True, ttl=1800)
    trafego_raw = cached_cu_tasks(TRAFEGO_LIST, include_closed=False, ttl=1800)
    tasks_raw = copy_raw
    tasks = [parse_cu_task(t) for t in tasks_raw if re.match(r"\s*\[", t["name"])]

    # Build TRÁFEGO index: map task names to their COPY counterparts via assignee
    # TRÁFEGO tasks often have the assignee as the copywriter who wrote the creative
    trafego_tasks = [parse_cu_task(t) for t in trafego_raw if re.match(r"\s*\[", t["name"])]
    tasks.extend(trafego_tasks)

    log(f"  {len(copy_raw)} COPY + {len(trafego_raw)} TRÁFEGO = {len(tasks)} processadas")

    # Build final
    log("Matching criativos → copywriters...")
    creatives = []
    for (base_id, version, nicho), data in creative_agg.items():
        cw, ed, mt = find_copywriter(base_id, version, nicho, tasks)
        # Fix C63 → AD63, REAPER
        if base_id == "C63" and nicho == "EM":
            base_id = "AD63"
            cw = "REAPER"
            mt = "MANUAL"
        rf = data["front_rev"] / data["cost"] if data["cost"] > 0 else 0
        rt = data["total_rev"] / data["cost"] if data["cost"] > 0 else 0
        creatives.append({
            "base_id": base_id, "version": version, "nicho": nicho,
            "copywriter": cw, "editor": ed,
            "cost": data["cost"], "front_rev": data["front_rev"],
            "total_rev": data["total_rev"], "mc_br": data["mc_br"],
            "roas_front": rf, "roas_total": rt,
            "vendas": data["vendas"], "match_type": mt,
        })

    creatives.sort(key=lambda x: x["vendas"], reverse=True)
    log(f"  {len(creatives)} criativos processados")
    return creatives


# ============================================================
# DOCX GENERATION
# ============================================================

def generate_docx(creatives, date_from, date_to, output_path=None):
    """Gera o .docx completo do relatório."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    def fmt_brl(v):
        if abs(v) >= 1000:
            s = f"{abs(v):,.0f}".replace(",", ".")
        else:
            s = f"{abs(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        return f"R$ {s}" if v >= 0 else f"-R$ {s}"

    def fmt_roas(v):
        return f"{v:.2f}"

    def fmt_pct(v):
        return f"{v:.1f}%"

    def fmt_int(v):
        return f"{v:,}".replace(",", ".")

    DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
    MEDIUM_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
    ALT_ROW_BG = "F2F6FA"
    GREEN = RGBColor(0x27, 0x7D, 0x3E)
    ORANGE = RGBColor(0xCC, 0x7A, 0x00)
    RED = RGBColor(0xC0, 0x39, 0x2B)
    BLACK = RGBColor(0x2D, 0x2D, 0x2D)
    GRAY = RGBColor(0x6B, 0x6B, 0x6B)

    def roas_color(v):
        if v >= 2.0: return GREEN
        if v >= 1.5: return ORANGE
        if v >= 1.0: return RGBColor(0xE6, 0x7E, 0x22)
        return RED

    def set_cell_bg(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_text(cell, text, bold=False, size=8, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ""
        p = cell.paragraphs[0]; p.alignment = align; p.space_before = Pt(1); p.space_after = Pt(1)
        run = p.add_run(str(text))
        run.font.size = Pt(size); run.font.color.rgb = color; run.font.name = "Calibri"; run.bold = bold
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def set_cell_number(cell, text, bold=False, size=8, color=BLACK):
        set_cell_text(cell, text, bold=bold, size=size, color=color, align=WD_ALIGN_PARAGRAPH.RIGHT)

    def add_header_row(table, headers):
        row = table.rows[0]
        for i, h in enumerate(headers):
            set_cell_text(row.cells[i], h, bold=True, size=7.5, color=RGBColor(0xFF, 0xFF, 0xFF))
            set_cell_bg(row.cells[i], "1B3A5C")
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Compute summaries
    total_cost = sum(c["cost"] for c in creatives)
    total_front = sum(c["front_rev"] for c in creatives)
    total_total = sum(c["total_rev"] for c in creatives)
    total_mc = sum(c["mc_br"] for c in creatives)
    total_vendas = sum(c["vendas"] for c in creatives)

    copy_summary = defaultdict(lambda: {"cost": 0, "front_rev": 0, "total_rev": 0, "mc_br": 0, "vendas": 0, "criativos": 0, "nichos": set()})
    for c in creatives:
        d = copy_summary[c["copywriter"]]
        d["cost"] += c["cost"]; d["front_rev"] += c["front_rev"]; d["total_rev"] += c["total_rev"]
        d["mc_br"] += c["mc_br"]; d["vendas"] += c["vendas"]; d["criativos"] += 1
        if c["nicho"]: d["nichos"].add(c["nicho"])
    sorted_copy = sorted(copy_summary.items(), key=lambda x: x[1]["vendas"], reverse=True)

    df_label = datetime.strptime(date_from, "%Y-%m-%d").strftime("%d/%m/%Y")
    dt_label = datetime.strptime(date_to, "%Y-%m-%d").strftime("%d/%m/%Y")

    # Build doc
    doc = Document()
    for section in doc.sections:
        section.orientation = 1; section.page_width = Cm(42.0); section.page_height = Cm(29.7)
        section.left_margin = Cm(1.5); section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10); style.font.color.rgb = BLACK

    # COVER
    doc.add_paragraph(""); doc.add_paragraph("")
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("RELATÓRIO DE PERFORMANCE POR CRIATIVO")
    run.font.size = Pt(26); run.font.color.rgb = DARK_BLUE; run.bold = True
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = st.add_run("Análise de Produtividade e Lucratividade por Copywriter")
    run.font.size = Pt(14); run.font.color.rgb = MEDIUM_BLUE
    mt = doc.add_paragraph(); mt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mt.add_run(f"Período: {df_label} a {dt_label}\nIMPERA PRODUTOS NATURAIS")
    run.font.size = Pt(11); run.font.color.rgb = GRAY
    doc.add_paragraph("")
    g = doc.add_paragraph(); g.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = g.add_run(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
    run.font.size = Pt(9); run.font.color.rgb = GRAY
    doc.add_page_break()

    # 1. RESUMO EXECUTIVO
    h = doc.add_heading("1. Resumo Executivo", level=1)
    for r in h.runs: r.font.color.rgb = DARK_BLUE
    p = doc.add_paragraph()
    p.add_run(f"Este relatório analisa a performance individual de criativos publicitários da IMPERA no período de {df_label} a {dt_label}, cruzando dados de performance do RedTrack com a atribuição de copywriters via ClickUp.").font.size = Pt(10)
    doc.add_paragraph("")
    tb = doc.add_table(rows=1, cols=6); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(tb, ["Indicador", "Custo Total", "Fat. Front", "Fat. Total", "MC Bruta", "Vendas"])
    r = tb.add_row()
    set_cell_text(r.cells[0], "Total Geral", bold=True, size=9)
    set_cell_number(r.cells[1], fmt_brl(total_cost), size=9, bold=True)
    set_cell_number(r.cells[2], fmt_brl(total_front), size=9, bold=True)
    set_cell_number(r.cells[3], fmt_brl(total_total), size=9, bold=True)
    set_cell_number(r.cells[4], fmt_brl(total_mc), size=9, bold=True, color=GREEN if total_mc > 0 else RED)
    set_cell_number(r.cells[5], fmt_int(total_vendas), size=9, bold=True)
    doc.add_paragraph("")
    h2 = doc.add_heading("Principais Achados", level=2)
    for r in h2.runs: r.font.color.rgb = MEDIUM_BLUE
    with_sales = len([c for c in creatives if c["vendas"] > 0])
    findings = [
        f"{len(creatives)} criativos únicos analisados, {with_sales} com pelo menos 1 venda.",
        f"Top copywriter: {sorted_copy[0][0]} com {fmt_int(sorted_copy[0][1]['vendas'])} vendas ({sorted_copy[0][1]['vendas']/total_vendas*100:.1f}%)." if total_vendas > 0 else "",
        f"Criativo #1: {creatives[0]['base_id']} {creatives[0]['version']} ({creatives[0]['copywriter']}) — {fmt_int(creatives[0]['vendas'])} vendas, ROAS Total {creatives[0]['roas_total']:.2f}." if creatives else "",
        f"ROAS Front médio: {total_front/total_cost:.2f} | ROAS Total médio: {total_total/total_cost:.2f}." if total_cost > 0 else "",
        f"MC bruta total: {fmt_brl(total_mc)} ({total_mc/total_total*100:.1f}% do fat. total)." if total_total > 0 else "",
    ]
    for f in findings:
        if f: doc.add_paragraph(f, style="List Bullet")
    doc.add_page_break()

    # 2. METODOLOGIA
    h = doc.add_heading("2. Observações e Metodologia", level=1)
    for r in h.runs: r.font.color.rgb = DARK_BLUE
    h2 = doc.add_heading("Fontes de Dados", level=2)
    for r in h2.runs: r.font.color.rgb = MEDIUM_BLUE
    for tt, desc in [("RedTrack", "Performance de campanhas e adgroups. Cópias unificadas automaticamente."), ("ClickUp", "Atribuição de copywriter via campo customizado nas listas COPY e EDIÇÃO.")]:
        p = doc.add_paragraph()
        run = p.add_run(f"{tt}: "); run.bold = True; run.font.size = Pt(10)
        p.add_run(desc).font.size = Pt(10)
    h2 = doc.add_heading("Campos Utilizados", level=2)
    for r in h2.runs: r.font.color.rgb = MEDIUM_BLUE
    fields = [("Custo", "cost"), ("Fat. Front", "revenuetype2 + revenuetype3"), ("ROAS Front", "Fat. Front ÷ Custo"), ("MC BR", "profit = total_revenue − cost"), ("Fat. Total", "total_revenue"), ("ROAS Total", "Fat. Total ÷ Custo"), ("Vendas", "convtype1 (Purchase total)")]
    tb = doc.add_table(rows=1, cols=2); add_header_row(tb, ["Campo", "Definição"])
    for fn, fd in fields:
        r = tb.add_row(); set_cell_text(r.cells[0], fn, bold=True, size=9); set_cell_text(r.cells[1], fd, size=9)
    h2 = doc.add_heading("Legenda ROAS", level=2)
    for r in h2.runs: r.font.color.rgb = MEDIUM_BLUE
    for val, label, color in [("≥ 2.00", "Saudável", GREEN), ("1.50 – 1.99", "Atenção", ORANGE), ("1.00 – 1.49", "Risco", RGBColor(0xE6, 0x7E, 0x22)), ("< 1.00", "Prejuízo", RED)]:
        p = doc.add_paragraph(); run = p.add_run(f"● {val} — {label}"); run.font.color.rgb = color; run.font.size = Pt(10); run.bold = True
    h2 = doc.add_heading("Cobertura de Atribuição", level=2)
    for r in h2.runs: r.font.color.rgb = MEDIUM_BLUE
    matched = sum(1 for c in creatives if c["match_type"] in ("MATCH", "VER_MATCH", "FIRST", "PREFIXO", "MANUAL"))
    c_inf = sum(1 for c in creatives if c["match_type"] == "C##_INFERRED")
    unm = sum(1 for c in creatives if c["match_type"] == "NONE")
    v_unm = sum(c["vendas"] for c in creatives if c["match_type"] == "NONE")
    for item in [f"Match direto: {matched} criativos", f"DOUGLAS* (C## s/ ClickUp): {c_inf} criativos", f"Sem atribuição: {unm} criativos ({v_unm} vendas — {v_unm/total_vendas*100:.1f}%)" if total_vendas > 0 else f"Sem atribuição: {unm}"]:
        doc.add_paragraph(item, style="List Bullet")
    p = doc.add_paragraph()
    run = p.add_run("Nota: "); run.bold = True; run.font.size = Pt(9); run.font.color.rgb = GRAY
    p.add_run("DOUGLAS* = atribuição presumida, pendente confirmação.").font.size = Pt(9)
    doc.add_page_break()

    # 3. ANÁLISE POR COPYWRITER
    h = doc.add_heading("3. Análise por Copywriter", level=1)
    for r in h.runs: r.font.color.rgb = DARK_BLUE
    for cw_idx, (copywriter, summary) in enumerate(sorted_copy):
        if cw_idx > 0: doc.add_paragraph("")
        rf = summary["front_rev"] / summary["cost"] if summary["cost"] > 0 else 0
        rt = summary["total_rev"] / summary["cost"] if summary["cost"] > 0 else 0
        pv = summary["vendas"] / total_vendas * 100 if total_vendas > 0 else 0
        cpa = summary["cost"] / summary["vendas"] if summary["vendas"] > 0 else 0
        h2 = doc.add_heading(f"3.{cw_idx+1}. {copywriter}", level=2)
        for r in h2.runs: r.font.color.rgb = MEDIUM_BLUE
        tb = doc.add_table(rows=1, cols=8); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_header_row(tb, ["Criativos", "Custo", "Fat. Front", "ROAS Front", "MC BR", "Fat. Total", "ROAS Total", "Vendas"])
        r = tb.add_row()
        set_cell_number(r.cells[0], fmt_int(summary["criativos"]), bold=True, size=9)
        set_cell_number(r.cells[1], fmt_brl(summary["cost"]), size=9)
        set_cell_number(r.cells[2], fmt_brl(summary["front_rev"]), size=9)
        set_cell_number(r.cells[3], fmt_roas(rf), size=9, bold=True, color=roas_color(rf))
        set_cell_number(r.cells[4], fmt_brl(summary["mc_br"]), size=9, color=GREEN if summary["mc_br"] > 0 else RED)
        set_cell_number(r.cells[5], fmt_brl(summary["total_rev"]), size=9)
        set_cell_number(r.cells[6], fmt_roas(rt), size=9, bold=True, color=roas_color(rt))
        set_cell_number(r.cells[7], fmt_int(summary["vendas"]), size=9, bold=True)
        p = doc.add_paragraph()
        run = p.add_run(f"Participação: {fmt_pct(pv)} das vendas | CPA médio: {fmt_brl(cpa)} | Nichos: {', '.join(sorted(summary['nichos']))}")
        run.font.size = Pt(9); run.font.color.rgb = GRAY

        # Nicho breakdown
        cw_cr = [c for c in creatives if c["copywriter"] == copywriter]
        nbd = defaultdict(lambda: {"cost": 0, "front_rev": 0, "total_rev": 0, "mc_br": 0, "vendas": 0, "criativos": 0})
        for c in cw_cr:
            d = nbd[c["nicho"] or "?"]; d["cost"] += c["cost"]; d["front_rev"] += c["front_rev"]
            d["total_rev"] += c["total_rev"]; d["mc_br"] += c["mc_br"]; d["vendas"] += c["vendas"]; d["criativos"] += 1
        sn = sorted(nbd.items(), key=lambda x: x[1]["vendas"], reverse=True)
        if len(sn) > 1 or (len(sn) == 1 and sn[0][0] != "?"):
            doc.add_heading("Breakdown por Nicho", level=3)
            tb = doc.add_table(rows=1, cols=8); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_header_row(tb, ["Nicho", "Criat.", "Custo", "Fat. Front", "ROAS F", "MC BR", "Fat. Total", "Vendas"])
            for ni, (nicho, nd) in enumerate(sn):
                r = tb.add_row()
                if ni % 2 == 0:
                    for cell in r.cells: set_cell_bg(cell, ALT_ROW_BG)
                nrf = nd["front_rev"] / nd["cost"] if nd["cost"] > 0 else 0
                set_cell_text(r.cells[0], NICHO_FULL.get(nicho, nicho), bold=True, size=8)
                set_cell_number(r.cells[1], fmt_int(nd["criativos"]), size=8)
                set_cell_number(r.cells[2], fmt_brl(nd["cost"]), size=8)
                set_cell_number(r.cells[3], fmt_brl(nd["front_rev"]), size=8)
                set_cell_number(r.cells[4], fmt_roas(nrf), size=8, color=roas_color(nrf), bold=True)
                set_cell_number(r.cells[5], fmt_brl(nd["mc_br"]), size=8, color=GREEN if nd["mc_br"] > 0 else RED)
                set_cell_number(r.cells[6], fmt_brl(nd["total_rev"]), size=8)
                set_cell_number(r.cells[7], fmt_int(nd["vendas"]), size=8, bold=True)

        # Top 10
        top_cw = sorted(cw_cr, key=lambda x: x["vendas"], reverse=True)[:10]
        doc.add_heading("Top 10 Criativos", level=3)
        tb = doc.add_table(rows=1, cols=9); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_header_row(tb, ["#", "Criativo", "Versão", "Nicho", "Custo", "Fat. Front", "ROAS F", "Fat. Total", "Vendas"])
        for ci, c in enumerate(top_cw):
            r = tb.add_row()
            if ci % 2 == 0:
                for cell in r.cells: set_cell_bg(cell, ALT_ROW_BG)
            set_cell_number(r.cells[0], str(ci + 1), size=8)
            set_cell_text(r.cells[1], c["base_id"], bold=True, size=8)
            set_cell_text(r.cells[2], c["version"], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(r.cells[3], NICHO_FULL.get(c["nicho"], "?")[:6], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_number(r.cells[4], fmt_brl(c["cost"]), size=8)
            set_cell_number(r.cells[5], fmt_brl(c["front_rev"]), size=8)
            set_cell_number(r.cells[6], fmt_roas(c["roas_front"]), size=8, color=roas_color(c["roas_front"]), bold=True)
            set_cell_number(r.cells[7], fmt_brl(c["total_rev"]), size=8)
            set_cell_number(r.cells[8], fmt_int(c["vendas"]), size=8, bold=True)

        # MC negativa
        neg = sorted([c for c in cw_cr if c["mc_br"] < 0 and c["cost"] > 50], key=lambda x: x["mc_br"])
        if neg:
            doc.add_heading("Criativos com MC Negativa", level=3)
            tb = doc.add_table(rows=1, cols=7)
            add_header_row(tb, ["Criativo", "Versão", "Nicho", "Custo", "MC BR", "ROAS F", "Vendas"])
            for ni, c in enumerate(neg[:10]):
                r = tb.add_row()
                if ni % 2 == 0:
                    for cell in r.cells: set_cell_bg(cell, ALT_ROW_BG)
                set_cell_text(r.cells[0], c["base_id"], bold=True, size=8)
                set_cell_text(r.cells[1], c["version"], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(r.cells[2], NICHO_FULL.get(c["nicho"], "?")[:6], size=8)
                set_cell_number(r.cells[3], fmt_brl(c["cost"]), size=8)
                set_cell_number(r.cells[4], fmt_brl(c["mc_br"]), size=8, color=RED, bold=True)
                set_cell_number(r.cells[5], fmt_roas(c["roas_front"]), size=8, color=roas_color(c["roas_front"]))
                set_cell_number(r.cells[6], fmt_int(c["vendas"]), size=8)
        if cw_idx < len(sorted_copy) - 1: doc.add_page_break()

    doc.add_page_break()

    # 4. VISÃO GLOBAL
    h = doc.add_heading("4. Visão Global Comparativa", level=1)
    for r in h.runs: r.font.color.rgb = DARK_BLUE
    h2 = doc.add_heading("Ranking por Copywriter", level=2)
    for r in h2.runs: r.font.color.rgb = MEDIUM_BLUE
    tb = doc.add_table(rows=1, cols=10); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(tb, ["#", "Copywriter", "Criat.", "Custo", "Fat. Front", "ROAS F", "MC BR", "Fat. Total", "ROAS T", "Vendas"])
    for ci, (cw, d) in enumerate(sorted_copy):
        r = tb.add_row()
        if ci % 2 == 0:
            for cell in r.cells: set_cell_bg(cell, ALT_ROW_BG)
        rf = d["front_rev"] / d["cost"] if d["cost"] > 0 else 0
        rt = d["total_rev"] / d["cost"] if d["cost"] > 0 else 0
        set_cell_number(r.cells[0], str(ci + 1), size=9)
        set_cell_text(r.cells[1], cw, bold=True, size=9)
        set_cell_number(r.cells[2], fmt_int(d["criativos"]), size=9)
        set_cell_number(r.cells[3], fmt_brl(d["cost"]), size=9)
        set_cell_number(r.cells[4], fmt_brl(d["front_rev"]), size=9)
        set_cell_number(r.cells[5], fmt_roas(rf), size=9, color=roas_color(rf), bold=True)
        set_cell_number(r.cells[6], fmt_brl(d["mc_br"]), size=9, color=GREEN if d["mc_br"] > 0 else RED)
        set_cell_number(r.cells[7], fmt_brl(d["total_rev"]), size=9)
        set_cell_number(r.cells[8], fmt_roas(rt), size=9, color=roas_color(rt), bold=True)
        set_cell_number(r.cells[9], fmt_int(d["vendas"]), size=9, bold=True)
    # Total
    r = tb.add_row()
    for cell in r.cells: set_cell_bg(cell, "1B3A5C")
    rft = total_front / total_cost if total_cost > 0 else 0
    rtt = total_total / total_cost if total_cost > 0 else 0
    vals = ["", "TOTAL", fmt_int(len(creatives)), fmt_brl(total_cost), fmt_brl(total_front), fmt_roas(rft), fmt_brl(total_mc), fmt_brl(total_total), fmt_roas(rtt), fmt_int(total_vendas)]
    for i, v in enumerate(vals):
        if i <= 1: set_cell_text(r.cells[i], v, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        else: set_cell_number(r.cells[i], v, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    doc.add_paragraph("")
    h2 = doc.add_heading("Participação e Eficiência", level=2)
    for r in h2.runs: r.font.color.rgb = MEDIUM_BLUE
    tb = doc.add_table(rows=1, cols=6); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(tb, ["Copywriter", "% Vendas", "% Fat.", "CPA", "ROAS Total", "MC/Venda"])
    for ci, (cw, d) in enumerate(sorted_copy):
        if d["cost"] < 100: continue
        r = tb.add_row()
        if ci % 2 == 0:
            for cell in r.cells: set_cell_bg(cell, ALT_ROW_BG)
        pv = d["vendas"] / total_vendas * 100 if total_vendas > 0 else 0
        pf = d["total_rev"] / total_total * 100 if total_total > 0 else 0
        cpa = d["cost"] / d["vendas"] if d["vendas"] > 0 else 0
        rt = d["total_rev"] / d["cost"] if d["cost"] > 0 else 0
        mcpv = d["mc_br"] / d["vendas"] if d["vendas"] > 0 else 0
        set_cell_text(r.cells[0], cw, bold=True, size=9)
        set_cell_number(r.cells[1], fmt_pct(pv), size=9, bold=True)
        set_cell_number(r.cells[2], fmt_pct(pf), size=9)
        set_cell_number(r.cells[3], fmt_brl(cpa), size=9)
        set_cell_number(r.cells[4], fmt_roas(rt), size=9, color=roas_color(rt), bold=True)
        set_cell_number(r.cells[5], fmt_brl(mcpv), size=9, color=GREEN if mcpv > 0 else RED)
    doc.add_page_break()

    # 5. BASE VS VARIAÇÕES
    h = doc.add_heading("5. Criativos Base vs Variações", level=1)
    for r in h.runs: r.font.color.rgb = DARK_BLUE
    p = doc.add_paragraph()
    p.add_run("Criativos agrupados pelo identificador base, com contagem de versões e performance agregada.").font.size = Pt(10)
    bg = defaultdict(lambda: {"versions": [], "total_cost": 0, "total_rev": 0, "total_vendas": 0, "copywriter": None, "nicho": None})
    for c in creatives:
        b = bg[c["base_id"]]; b["versions"].append(c); b["total_cost"] += c["cost"]; b["total_rev"] += c["total_rev"]; b["total_vendas"] += c["vendas"]
        if not b["copywriter"]: b["copywriter"] = c["copywriter"]
        if not b["nicho"]: b["nicho"] = c["nicho"]
    mv = sorted([(bid, g) for bid, g in bg.items() if len(g["versions"]) >= 2 and g["total_vendas"] >= 10], key=lambda x: -x[1]["total_vendas"])
    tb = doc.add_table(rows=1, cols=9); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(tb, ["Base", "Copy", "Nicho", "Versões", "Custo Total", "Fat. Total", "ROAS T", "Vendas", "Melhor Versão"])
    for ci, (bid, g) in enumerate(mv[:30]):
        r = tb.add_row()
        if ci % 2 == 0:
            for cell in r.cells: set_cell_bg(cell, ALT_ROW_BG)
        rt = g["total_rev"] / g["total_cost"] if g["total_cost"] > 0 else 0
        best = max(g["versions"], key=lambda x: x["vendas"])
        bl = best["version"] if best["version"] != "-" else "Base"
        set_cell_text(r.cells[0], bid, bold=True, size=8)
        set_cell_text(r.cells[1], g["copywriter"][:10], size=8)
        set_cell_text(r.cells[2], NICHO_FULL.get(g["nicho"], "?")[:6], size=8)
        set_cell_number(r.cells[3], str(len(g["versions"])), size=8, bold=True)
        set_cell_number(r.cells[4], fmt_brl(g["total_cost"]), size=8)
        set_cell_number(r.cells[5], fmt_brl(g["total_rev"]), size=8)
        set_cell_number(r.cells[6], fmt_roas(rt), size=8, color=roas_color(rt), bold=True)
        set_cell_number(r.cells[7], fmt_int(g["total_vendas"]), size=8, bold=True)
        set_cell_text(r.cells[8], f"{bl} ({fmt_int(best['vendas'])}v)", size=8)
    doc.add_page_break()

    # 6. RANKING COMPLETO
    h = doc.add_heading("6. Ranking Completo de Criativos", level=1)
    for r in h.runs: r.font.color.rgb = DARK_BLUE
    p = doc.add_paragraph()
    p.add_run(f"Todos os {len(creatives)} criativos, ordenados por vendas.").font.size = Pt(9)
    PAGE_SIZE = 80
    pages = [creatives[i:i + PAGE_SIZE] for i in range(0, len(creatives), PAGE_SIZE)]
    for pi, pc in enumerate(pages):
        if pi > 0:
            doc.add_page_break()
            p = doc.add_paragraph()
            run = p.add_run(f"Ranking Completo ({pi+1}/{len(pages)})")
            run.font.size = Pt(10); run.font.color.rgb = MEDIUM_BLUE; run.bold = True
        tb = doc.add_table(rows=1, cols=11); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_header_row(tb, ["#", "Criativo", "Versão", "Nicho", "Copy", "Custo", "Fat.Front", "ROAS F", "MC BR", "Fat.Total", "Vendas"])
        start = pi * PAGE_SIZE
        for ci, c in enumerate(pc):
            r = tb.add_row()
            if ci % 2 == 0:
                for cell in r.cells: set_cell_bg(cell, ALT_ROW_BG)
            set_cell_number(r.cells[0], str(start + ci + 1), size=7)
            set_cell_text(r.cells[1], c["base_id"], bold=True, size=7)
            set_cell_text(r.cells[2], c["version"], size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(r.cells[3], c["nicho"] or "?", size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(r.cells[4], c["copywriter"][:12], size=7)
            set_cell_number(r.cells[5], fmt_brl(c["cost"]), size=7)
            set_cell_number(r.cells[6], fmt_brl(c["front_rev"]), size=7)
            set_cell_number(r.cells[7], fmt_roas(c["roas_front"]), size=7, color=roas_color(c["roas_front"]), bold=True)
            set_cell_number(r.cells[8], fmt_brl(c["mc_br"]), size=7, color=GREEN if c["mc_br"] > 0 else RED)
            set_cell_number(r.cells[9], fmt_brl(c["total_rev"]), size=7)
            set_cell_number(r.cells[10], fmt_int(c["vendas"]), size=7, bold=True)

    # Signature
    doc.add_paragraph("")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("─" * 40).font.color.rgb = GRAY
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("GPDR — Iago Almeida, assistido por Claude")
    run.font.size = Pt(9); run.font.color.rgb = GRAY; run.italic = True

    # Save
    if not output_path:
        df_short = datetime.strptime(date_from, "%Y-%m-%d").strftime("%d%m")
        dt_short = datetime.strptime(date_to, "%Y-%m-%d").strftime("%d%m")
        output_path = os.path.expanduser(f"~/Documents/Relatorio_Performance_Criativos_{df_short}a{dt_short}.docx")
    doc.save(output_path)
    return output_path


# ============================================================
# TELEGRAM: send document
# ============================================================

def send_telegram_document(file_path, caption=""):
    if not TELEGRAM_TOKEN or not TELEGRAM_CHAT_ID:
        print("Telegram não configurado.")
        return
    url = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendDocument"
    import mimetypes
    boundary = "----FormBoundary7MA4YWxkTrZu0gW"
    filename = os.path.basename(file_path)
    with open(file_path, "rb") as f:
        file_data = f.read()
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="chat_id"\r\n\r\n{TELEGRAM_CHAT_ID}\r\n'
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="caption"\r\n\r\n{caption}\r\n'
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="document"; filename="{filename}"\r\n'
        f"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
    ).encode() + file_data + f"\r\n--{boundary}--\r\n".encode()
    req = urllib.request.Request(url, data=body, method="POST")
    req.add_header("Content-Type", f"multipart/form-data; boundary={boundary}")
    urllib.request.urlopen(req)
    print(f"Documento enviado via Telegram: {filename}")


def send_telegram_msg(text):
    if not TELEGRAM_TOKEN or not TELEGRAM_CHAT_ID:
        return
    url = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendMessage"
    body = json.dumps({"chat_id": TELEGRAM_CHAT_ID, "text": text}).encode()
    req = urllib.request.Request(url, data=body, method="POST")
    req.add_header("Content-Type", "application/json")
    urllib.request.urlopen(req)


# ============================================================
# PERIOD HELPERS
# ============================================================

def get_last_week():
    """Retorna segunda anterior 00:00 até domingo anterior 23:59."""
    today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    # Domingo anterior
    days_since_sunday = today.weekday() + 1  # Monday=0 → +1=1
    if days_since_sunday == 7:
        days_since_sunday = 0  # Se hoje é domingo, pega a semana anterior
    last_sunday = today - timedelta(days=days_since_sunday)
    last_monday = last_sunday - timedelta(days=6)
    return last_monday.strftime("%Y-%m-%d"), last_sunday.strftime("%Y-%m-%d")


def parse_date_arg(arg):
    """Parse DD/MM or DD/MM/YYYY."""
    arg = arg.strip()
    if re.match(r"^\d{2}/\d{2}/\d{4}$", arg):
        return datetime.strptime(arg, "%d/%m/%Y").strftime("%Y-%m-%d")
    elif re.match(r"^\d{2}/\d{2}$", arg):
        year = datetime.now().year
        return datetime.strptime(f"{arg}/{year}", "%d/%m/%Y").strftime("%Y-%m-%d")
    return None


# ============================================================
# MAIN
# ============================================================

def run(date_from, date_to, notify=False):
    """Executa coleta + geração + envio opcional."""
    df_label = datetime.strptime(date_from, "%Y-%m-%d").strftime("%d/%m")
    dt_label = datetime.strptime(date_to, "%Y-%m-%d").strftime("%d/%m")
    print(f"Período: {df_label} a {dt_label}")

    if notify:
        send_telegram_msg(f"Gerando Relatório de Performance por Criativo ({df_label} a {dt_label})...")

    creatives = collect_data(date_from, date_to)
    output = generate_docx(creatives, date_from, date_to)
    print(f"Relatório salvo: {output}")

    if notify:
        total_vendas = sum(c["vendas"] for c in creatives)
        total_rev = sum(c["total_rev"] for c in creatives)
        caption = f"Relatório Performance Criativos | {df_label} a {dt_label}\n{len(creatives)} criativos | {total_vendas:,} vendas | R${total_rev:,.0f} fat. total".replace(",", ".")
        send_telegram_document(output, caption)

    return output


if __name__ == "__main__":
    args = [a for a in sys.argv[1:] if a != "--notify"]
    notify = "--notify" in sys.argv

    if len(args) >= 2:
        d1 = parse_date_arg(args[0])
        d2 = parse_date_arg(args[1])
        if not d1 or not d2:
            print("Formato: DD/MM ou DD/MM/YYYY")
            sys.exit(1)
        run(d1, d2, notify=notify)
    else:
        d1, d2 = get_last_week()
        run(d1, d2, notify=notify)
